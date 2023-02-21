import tkinter as tk
import docx
from tkinter import messagebox
from tkinter import filedialog
import pyttsx3
from tkinter import Tk, Text, DISABLED
from tkinter import *
from tkinter import ttk
# which is used to save file in any extension
from tkinter.filedialog import asksaveasfile
import os
import tkinter.font as font
from PIL import ImageTk, Image

global counter
counter = 1

def translate_to_braille(text):

    braille_dict = {
  'a': '⠁',
  'b': '⠃',
  'c': '⠉',
  'd': '⠙',
  'e': '⠑',
  'f': '⠋',
  'g': '⠛',
  'h': '⠓',
  'i': '⠊',
  'j': '⠚',
  'k': '⠅',
  'l': '⠇',
  'm': '⠍',
  'n': '⠝',
  'o': '⠕',
  'p': '⠏',
  'q': '⠟',
  'r': '⠗',
  's': '⠎',
  't': '⠞',
  'u': '⠥',
  'v': '⠧',
  'w': '⠺',
  'x': '⠭',
  'y': '⠽',
  'z': '⠵',
  ' ': '  ',
   '1': '⠧'

    }
    braille = ''
    for char in text:
        if char.lower() in braille_dict:
            braille += braille_dict[char.lower()] + ' '
        else:
            braille += char
    return braille


def upload_file():
    global counter

    if counter < 2:
        new_window = tk.Tk()
        new_window.title('Upload Function')
        new_window.geometry("600x100")
        new_window.configure(bg='#fce388')

    counter +=1

    def brFw ():
        top = Toplevel()
        top.title("BRF Braille Translation")
        top.geometry("1100x700")
        top.configure(bg='#fce388')

        upload_textbox = Text(top, height=28, width=67)
        upload_textbox.grid(row=0, column=0, padx=5, pady=5, sticky="nsew",)
        upload_textbox.delete('1.0', END)


        braille_label = Text(top, height= 28, width=67)
        braille_label.grid(row=0, column=1, padx=5, pady=5, sticky="nsew")
        braille_label.delete('1.0', END)
        braille_label.configure(state='disabled')
        upload_textbox.configure(state='disabled')

        def on_upload_textbox_scroll(event):
            yview = upload_textbox.yview()


            braille_label.yview_moveto(yview[0])
        upload_textbox.bind("<MouseWheel>", on_upload_textbox_scroll)

        def on_braille_label_scroll(event):

            yview = braille_label.yview()


            upload_textbox.yview_moveto(yview[0])
        braille_label.bind("<MouseWheel>", on_braille_label_scroll)

        def update_page_counter(event=None):
            words = upload_textbox.get("1.0", "end-1c").split()
            total_pages = len(words)//250
            yview = upload_textbox.yview()
            current_page = int((yview[0] + yview[1])*total_pages)
            page_label.config(text="Page: " + str(current_page) +" of "+str(total_pages))

        page_label =tk.Label(top, text="Page: 1 of 1")
        page_label.grid(row=0, column=3, padx=5, pady=5, sticky="w")

        upload_textbox.bind("<MouseWheel>", update_page_counter,on_upload_textbox_scroll)

        def text_speech():
            engine = pyttsx3.init()
            engine.say(upload_textbox.get('1.0', 'END'))
            engine.runAndWait()
            upload_textbox.delete(0, END)

        speak = tk.Button(top, text="Speak", bd = '5', bg='#f38181', command=text_speech)
        speak.grid(row=2, column=0, padx=5, pady=5, sticky="nsew")

        def save():
            files = [('All Files', '*.*'),
            ('Text Document', '*.txt'),
            ('Braille Ready Format', '*.brf')]
            f=asksaveasfile(mode='w', filetypes = files, defaultextension = files)
            save_file=str(upload_textbox.get(0.1, END))
            f.write(save_file)
            f.close

        saveButton = tk.Button(top, text="Save File", bd = '5', bg='#f38181', command=save)
        saveButton.grid(row=1, column=1, padx=5, pady=5, sticky="nsew")

        print = tk.Button(top, text="Print", bd = '5', bg='#f38181', command=save)
        print.grid(row=2, column=1, padx=5, pady=5, sticky="nsew")

        def clear_text():
            upload_textbox.configure(state='normal')
            upload_textbox.delete(0.1, END)
            braille_label.configure(state='normal')
            braille_label.delete(0.1, END)

        def translate_text(text):
            braille = translate_to_braille(text)
            braille_label.insert(1.0, braille)
            braille_label.configure(state='disabled')


        def open_file():

            filepath = filedialog.askopenfilename(filetypes=[("Word files", "*.docx"), ("text files","*.txt"), ("BRF files", "*.brf")])
            if not filepath:
                return
            _, file_extension = os.path.splitext(filepath)
            if file_extension == '.docx':
                doc = docx.Document(filepath)
                text = '\n'.join([para.text for para in doc.paragraphs])
            elif file_extension == '.txt':
                with open(filepath, 'r') as f:
                    text = f.read()
            elif file_extension == '.brf':
                with open(filepath, 'r') as f:
                    text = f.read()

            braille_label.configure(state='normal')
            upload_textbox.configure(state='normal')
            upload_textbox.delete('1.0', END)
            braille_label.delete('1.0', END)
            translate_text(text)
            upload_textbox.insert("1.0", text)
            upload_textbox.configure(state='disabled')

        open_button = tk.Button(top, text="OPEN AND TRANSLATE FILE", bd = '5', bg='#f38181', command= lambda: [clear_text(), open_file()])
        open_button.grid(row=1, column=0, padx=5, pady=5, sticky="nsew")

        Grid.rowconfigure(top, 0, weight=8)
        Grid.rowconfigure(top, 1, weight=1)
        Grid.rowconfigure(top, 2, weight=1)
        Grid.columnconfigure(top, 0, weight=1)
        Grid.columnconfigure(top, 1, weight=1)

        top.mainloop()


    def txTw ():
        top1 = Toplevel()
        top1.title("TXT Braille Translation")
        top1.geometry("1100x700")
        top1.configure(bg='#fce388')

        upload_textbox = Text(top1, height=28, width=67)
        upload_textbox.grid(row=0, column=0, padx=5, pady=5, sticky="nsew",)
        upload_textbox.delete('1.0', END)


        braille_label = Text(top1, height= 28, width=67)
        braille_label.grid(row=0, column=1, padx=5, pady=5, sticky="nsew")
        braille_label.delete('1.0', END)
        braille_label.configure(state='disabled')
        upload_textbox.configure(state='disabled')

        def on_upload_textbox_scroll(event):
            yview = upload_textbox.yview()


            braille_label.yview_moveto(yview[0])
        upload_textbox.bind("<MouseWheel>", on_upload_textbox_scroll)

        def on_braille_label_scroll(event):

            yview = braille_label.yview()


            upload_textbox.yview_moveto(yview[0])
        braille_label.bind("<MouseWheel>", on_braille_label_scroll)

        def update_page_counter(event=None):
            words = upload_textbox.get("1.0", "end-1c").split()
            total_pages = len(words)//250
            yview = upload_textbox.yview()
            current_page = int((yview[0] + yview[1])*total_pages)
            page_label.config(text="Page: " + str(current_page) +" of "+str(total_pages))

        page_label =tk.Label(top1, text="Page: 1 of 1")
        page_label.grid(row=0, column=3, padx=5, pady=5, sticky="w")

        upload_textbox.bind("<MouseWheel>", update_page_counter,on_upload_textbox_scroll)

        def text_speech():
            engine = pyttsx3.init()
            engine.say(upload_textbox.get('1.0', 'END'))
            engine.runAndWait()
            upload_textbox.delete(0, END)

        speak = tk.Button(top1, text="Speak", bd = '5', bg='#f38181', command=text_speech)
        speak.grid(row=2, column=0, padx=5, pady=5, sticky="nsew")

        def save():
            files = [('All Files', '*.*'),
            ('Text Document', '*.txt'),
            ('Braille Ready Format', '*.brf')]
            f=asksaveasfile(mode='w', filetypes = files, defaultextension = files)
            save_file=str(upload_textbox.get(0.1, END))
            f.write(save_file)
            f.close

        saveButton = tk.Button(top1, text="Save File", bd = '5', bg='#f38181', command=save)
        saveButton.grid(row=1, column=1, padx=5, pady=5, sticky="nsew")

        print = tk.Button(top1, text="Print", bd = '5', bg='#f38181', command=save)
        print.grid(row=2, column=1, padx=5, pady=5, sticky="nsew")

        def clear_text():
            upload_textbox.configure(state='normal')
            upload_textbox.delete(0.1, END)
            braille_label.configure(state='normal')
            braille_label.delete(0.1, END)

        def translate_text(text):
            braille = translate_to_braille(text)
            braille_label.insert(1.0, braille)
            braille_label.configure(state='disabled')


        def open_file():

            filepath = filedialog.askopenfilename(filetypes=[("Word files", "*.docx"), ("text files","*.txt"), ("BRF files", "*.brf")])
            if not filepath:
                return
            _, file_extension = os.path.splitext(filepath)
            if file_extension == '.docx':
                doc = docx.Document(filepath)
                text = '\n'.join([para.text for para in doc.paragraphs])
            elif file_extension == '.txt':
                with open(filepath, 'r') as f:
                    text = f.read()
            elif file_extension == '.brf':
                with open(filepath, 'r') as f:
                    text = f.read()

            braille_label.configure(state='normal')
            upload_textbox.configure(state='normal')
            upload_textbox.delete('1.0', END)
            braille_label.delete('1.0', END)
            translate_text(text)
            upload_textbox.insert("1.0", text)
            upload_textbox.configure(state='disabled')

        open_button = tk.Button(top1, text="OPEN AND TRANSLATE FILE", bd = '5', bg='#f38181', command= lambda: [clear_text(), open_file()])
        open_button.grid(row=1, column=0, padx=5, pady=5, sticky="nsew")

        Grid.rowconfigure(top1, 0, weight=8)
        Grid.rowconfigure(top1, 1, weight=1)
        Grid.rowconfigure(top1, 2, weight=1)
        Grid.columnconfigure(top1, 0, weight=1)
        Grid.columnconfigure(top1, 1, weight=1)


        top1.mainloop()

    def doCw ():
        top2 = Toplevel()
        top2.title("TXT Braille Translation")
        top2.geometry("1100x700")
        top2.configure(bg='#fce388')

        upload_textbox = Text(top2, height=28, width=67)
        upload_textbox.grid(row=0, column=0, padx=5, pady=5, sticky="nsew",)
        upload_textbox.delete('1.0', END)


        braille_label = Text(top2, height= 28, width=67)
        braille_label.grid(row=0, column=1, padx=5, pady=5, sticky="nsew")
        braille_label.delete('1.0', END)
        braille_label.configure(state='disabled')
        upload_textbox.configure(state='disabled')

        def on_upload_textbox_scroll(event):
            yview = upload_textbox.yview()


            braille_label.yview_moveto(yview[0])
        upload_textbox.bind("<MouseWheel>", on_upload_textbox_scroll)

        def on_braille_label_scroll(event):

            yview = braille_label.yview()


            upload_textbox.yview_moveto(yview[0])
        braille_label.bind("<MouseWheel>", on_braille_label_scroll)

        def update_page_counter(event=None):
            words = upload_textbox.get("1.0", "end-1c").split()
            total_pages = len(words)//250
            yview = upload_textbox.yview()
            current_page = int((yview[0] + yview[1])*total_pages)
            page_label.config(text="Page: " + str(current_page) +" of "+str(total_pages))

        page_label =tk.Label(top2, text="Page: 1 of 1")
        page_label.grid(row=0, column=3, padx=5, pady=5, sticky="w")

        upload_textbox.bind("<MouseWheel>", update_page_counter,on_upload_textbox_scroll)

        def text_speech():
            engine = pyttsx3.init()
            engine.say(upload_textbox.get('1.0', 'END'))
            engine.runAndWait()
            upload_textbox.delete(0, END)

        speak = tk.Button(top2, text="Speak", bd = '5', bg='#f38181', command=text_speech)
        speak.grid(row=2, column=0, padx=5, pady=5, sticky="nsew")

        def save():
            files = [('All Files', '*.*'),
            ('Text Document', '*.txt'),
            ('Braille Ready Format', '*.brf')]
            f=asksaveasfile(mode='w', filetypes = files, defaultextension = files)
            save_file=str(upload_textbox.get(0.1, END))
            f.write(save_file)
            f.close

        saveButton = tk.Button(top2, text="Save File", bd = '5', bg='#f38181', command=save)
        saveButton.grid(row=1, column=1, padx=5, pady=5, sticky="nsew")

        print = tk.Button(top2, text="Print", bd = '5', bg='#f38181', command=save)
        print.grid(row=2, column=1, padx=5, pady=5, sticky="nsew")

        def clear_text():
            upload_textbox.configure(state='normal')
            upload_textbox.delete(0.1, END)
            braille_label.configure(state='normal')
            braille_label.delete(0.1, END)

        def translate_text(text):
            braille = translate_to_braille(text)
            braille_label.insert(1.0, braille)
            braille_label.configure(state='disabled')


        def open_file():

            filepath = filedialog.askopenfilename(filetypes=[("Word files", "*.docx"), ("text files","*.txt"), ("BRF files", "*.brf")])
            if not filepath:
                return
            _, file_extension = os.path.splitext(filepath)
            if file_extension == '.docx':
                doc = docx.Document(filepath)
                text = '\n'.join([para.text for para in doc.paragraphs])
            elif file_extension == '.txt':
                with open(filepath, 'r') as f:
                    text = f.read()
            elif file_extension == '.brf':
                with open(filepath, 'r') as f:
                    text = f.read()

            braille_label.configure(state='normal')
            upload_textbox.configure(state='normal')
            upload_textbox.delete('1.0', END)
            braille_label.delete('1.0', END)
            translate_text(text)
            upload_textbox.insert("1.0", text)
            upload_textbox.configure(state='disabled')

        open_button = tk.Button(top2, text="OPEN AND TRANSLATE FILE", bd = '5', bg='#f38181', command= lambda: [clear_text(), open_file()])
        open_button.grid(row=1, column=0, padx=5, pady=5, sticky="nsew")

        Grid.rowconfigure(top2, 0, weight=8)
        Grid.rowconfigure(top2, 1, weight=1)
        Grid.rowconfigure(top2, 2, weight=1)
        Grid.columnconfigure(top2, 0, weight=1)
        Grid.columnconfigure(top2, 1, weight=1)

        top2.mainloop()

    BRF = tk.Button(new_window, text="BRF", bd = '5', bg='#f38181', command=brFw)
    BRF.grid(row=0, column=0, padx=5, pady=5, sticky="nsew")
    TXT = tk.Button(new_window, text="TXT", bd = '5', bg='#f38181', command=txTw)
    TXT.grid(row=0, column=1, padx=5, pady=5, sticky="nsew")
    DOC = tk.Button(new_window, text="DOC", bd = '5', bg='#f38181', command=doCw)
    DOC.grid(row=0, column=2, padx=5, pady=5, sticky="nsew")

    Grid.rowconfigure(new_window, 0, weight=1)
    Grid.columnconfigure(new_window, 0, weight=1)
    Grid.columnconfigure(new_window, 1, weight=1)
    Grid.columnconfigure(new_window, 2, weight=1)


def show_text_input():

    global counter

    if counter < 2:

        # Create a new window for text input
        new_window = tk.Tk()
        new_window.title('Text Input')
        new_window.geometry("1100x700")
        new_window.configure(bg='#fce388')
    counter +=1

    initial_textbox = Text(new_window, height=15, width=67)
    initial_textbox.grid(row=0, column=0, padx=5, pady=5, sticky="nsew")
    translated_textbox = Text(new_window, height=15, width=67)
    translated_textbox.grid(row=0, column=2, padx=5, pady=5, sticky="nsew")
    translated_textbox.configure(state='disabled')

    def translate():
        # Get text from the input widget and translate it to braille
        text = initial_textbox.get('1.0', 'end')
        braille = translate_to_braille(text)
        # Display braille translation
        output_label = translated_textbox.insert(0.1, braille)
        translated_textbox.configure(state='disable')


    def text_speech():
        engine = pyttsx3.init()
        engine.say(initial_textbox.get('1.0', 'end'))
        engine.runAndWait()

    def clear_text():
        translated_textbox.configure(state='normal')
        translated_textbox.delete(0.1, END)

    def save1():
        files = [('All Files', '*.*'),
        ('Text Document', '*.txt'),
        ('Braille Ready Format', '*.brf')]
        f=asksaveasfile(mode='w', filetypes = files, defaultextension = files)
        save_file=str(initial_textbox.get(0.1, END))
        f.write(save_file)
        f.close

    translate_button = tk.Button(new_window, text='Translate', width = 20, height = 2, bd = '5', bg='#f38181', command= lambda: [clear_text(),translate()])
    translate_button.grid(row=0, column=1, padx=5, pady=300,  sticky="n")
    saveButton = tk.Button(new_window, text="Save File", bg='#f38181', command=save1)
    saveButton.grid(row=1, column=2, padx=5, pady=5,  sticky="nsew")
    Print = tk.Button(new_window, text = "Print", bd = '5', bg='#f38181', command = new_window.destroy)
    Print.grid(row=2, column=2, padx=5, pady=5, sticky="nsew")
    speak = tk.Button(new_window, text="Speak", bd = '5', bg='#f38181', command=text_speech)
    speak.grid(row=1, column=0, padx=5, pady=5, sticky="nsew")
    idPrint = tk.Button(new_window, text="Print Directly", bd = '5', bg='#f38181', command=save1)
    idPrint.grid(row=2, column=0, padx=5, pady=5, sticky="nsew")

    l2=tk.Label(new_window,text=0,font=30, width=20)
    l2.grid(row=0,column=0,padx=5,pady=5, sticky ="se")



    def my_upd(value):
        my_str=initial_textbox.get('1.0','end-1c') #The input string except the last line break
        char_numbers = len(my_str.split())
        l2.config(text=str(char_numbers)) # display number of chars

        if(char_numbers > 399):
            initial_textbox.delete('end-2c') # remove last char of text widget

    initial_textbox.bind('<KeyRelease>',my_upd) # Key release event to call function.

    Grid.rowconfigure(new_window, 0, weight=6)
    Grid.rowconfigure(new_window, 1, weight=2)
    Grid.rowconfigure(new_window, 2, weight=2)
    Grid.columnconfigure(new_window, 0, weight=4)
    Grid.columnconfigure(new_window, 1, weight=0)
    Grid.columnconfigure(new_window, 2, weight=4)

# Create main window
root = tk.Tk()
root.title('Braille Translator')
root.geometry("700x700")
root.configure(bg='#fce388')

#Logo

Logo = Image.open("C:\\Users\\Calculated\\Downloads\\lloyd\\123.png")
Logo = Logo.resize((110, 100))
render = ImageTk.PhotoImage(Logo,)
Logo = tk.Label(root, image=render)
Logo.grid(row=0, column=1, sticky = "s")
root.iconphoto(False, render)
Logo.configure(bg='#fce388')



Title = Label(root, text = "CALI-BRAILLE")
T_font = font.Font(family='Times New Roman', size=30, weight="bold")
Title['font'] = T_font
Title.grid(row=1, column=1, sticky = "n")
Title.configure(bg='#fce388')

Subtitle1 = Label(root, text = "AN AUTOMATED BRAILLE PRINTER WITH TRANSLATION")
S_font1 = font.Font(family='Times New Roman', size=15, weight="bold")
Subtitle1['font'] = S_font1
Subtitle1.grid(row=2, column=1, sticky = "n")
Subtitle1.configure(bg='#fce388')

Subtitle2 = Label(root, text = "SYSTEM AND TEXT-TO-SPEECH TECHNOLOGY")
Subtitle2['font'] = S_font1
Subtitle2.grid(row=3, column=1, sticky = "n")
Subtitle2.configure(bg='#fce388')


# Create buttons
B_font = font.Font(family='Montserrat', size=20, weight="bold")
B_font2 = font.Font(family='Montserrat', size=10, weight="bold")


upload_button = tk.Button(root, text='UPLOAD FILE', bd = '5', bg='#f38181', command= upload_file)
upload_button.grid(row=4, column=1, padx=5, pady=5,  sticky="nsew")
upload_button['font'] = B_font
text_input_button = tk.Button(root, text='TEXT INPUT', bd = '5', bg='#f38181', command=show_text_input)
text_input_button.grid(row=5, column=1, padx=5, pady=5,  sticky="nsew")
text_input_button['font'] = B_font

Help = tk.Button(root, text='HELP', width= 15, bd = '5', bg='#f38181', command=show_text_input)
Help.grid(row=7, column=1, pady=5,  sticky="e")
About_us = tk.Button(root, text='ABOUT US', width= 15, bd = '5', bg='#f38181', command=show_text_input)
About_us.grid(row=7, column=1, padx=145, pady=5, sticky="e")
Help['font'] = B_font2
About_us['font'] = B_font2


Grid.rowconfigure(root, 0, weight=1)
Grid.rowconfigure(root, 1, weight=0)
Grid.rowconfigure(root, 2, weight=0)
Grid.rowconfigure(root, 3, weight=1)
Grid.rowconfigure(root, 4, weight=1)
Grid.rowconfigure(root, 5, weight=1)
Grid.rowconfigure(root, 6, weight=1)
Grid.rowconfigure(root, 7, weight=1)
Grid.rowconfigure(root, 8, weight=1)

Grid.columnconfigure(root, 0, weight=1)
Grid.columnconfigure(root, 1, weight=3)
Grid.columnconfigure(root, 2, weight=1)


#558y
# Run main loop
root.mainloop()
